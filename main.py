"""FastAPI backend for Matchwise."""

from __future__ import annotations

import json
import os
import re
from io import BytesIO
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Request, UploadFile
from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from pypdf import PdfReader

from prompts.application_prompt import APPLICATION_PACK_SCHEMA, build_application_prompt
from prompts.resume_prompt import RESUME_SCHEMA, build_resume_prompt
from scripts.ats_score import analyze


BASE_DIR = Path(__file__).resolve().parent
PUBLIC_DIR = BASE_DIR / "public"
MAX_TEXT_LENGTH = 45_000
MAX_FILE_SIZE = 8 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}

load_dotenv(BASE_DIR / ".env")

app = FastAPI(
    title="Matchwise API",
    description="Python API for ATS scoring, resume rewriting, and application outreach.",
    version="2.0.0",
    docs_url="/api/docs",
    redoc_url=None,
)


class ResumeInputs(BaseModel):
    jobDescription: str = ""
    currentResume: str = ""


class ApiError(Exception):
    def __init__(self, message: str, status_code: int = 500) -> None:
        super().__init__(message)
        self.message = message
        self.status_code = status_code


@app.exception_handler(ApiError)
async def handle_api_error(_request: Request, error: ApiError) -> JSONResponse:
    return JSONResponse(status_code=error.status_code, content={"error": error.message})


@app.exception_handler(RequestValidationError)
async def handle_validation_error(_request: Request, _error: RequestValidationError) -> JSONResponse:
    return JSONResponse(status_code=400, content={"error": "The request body is invalid."})


@app.exception_handler(Exception)
async def handle_unexpected_error(_request: Request, error: Exception) -> JSONResponse:
    print(f"Unexpected server error: {error}")
    return JSONResponse(
        status_code=500,
        content={"error": "Could not generate the resume. Please try again."},
    )


@app.get("/api/health")
async def health() -> dict[str, bool]:
    return {"ok": True, "configured": bool(os.getenv("GEMINI_API_KEY"))}


@app.post("/api/extract-resume")
async def extract_resume(resume: UploadFile | None = File(default=None)) -> dict[str, Any]:
    if resume is None or not resume.filename:
        raise ApiError("Choose a PDF, DOCX, or TXT resume first.", 400)

    suffix = Path(resume.filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS and resume.content_type not in ALLOWED_MIME_TYPES:
        raise ApiError("Upload a PDF, DOCX, or TXT resume.", 415)

    content = await resume.read(MAX_FILE_SIZE + 1)
    await resume.close()
    if len(content) > MAX_FILE_SIZE:
        raise ApiError("The file is too large. Upload a resume smaller than 8 MB.", 413)

    try:
        text = extract_resume_text(content, suffix, resume.content_type or "")
    except ApiError:
        raise
    except Exception as error:
        raise ApiError(
            "Could not read this file. Try another PDF, DOCX, or TXT resume.", 422
        ) from error

    if not text.strip():
        raise ApiError(
            "No readable text was found in this file. Try pasting the resume text instead.",
            422,
        )

    cleaned = text.strip()
    return {"text": cleaned[:MAX_TEXT_LENGTH], "truncated": len(cleaned) > MAX_TEXT_LENGTH}


@app.post("/api/score")
async def score_resume(payload: ResumeInputs) -> dict[str, Any]:
    job_description, current_resume = validate_inputs(
        payload, "Add both the current resume and job description before scoring."
    )
    return {"analysis": analyze(job_description, current_resume)}


@app.post("/api/rewrite")
async def rewrite_resume(payload: ResumeInputs) -> dict[str, Any]:
    job_description, current_resume = validate_inputs(
        payload, "Add both the job description and current resume before generating."
    )
    ensure_gemini_key()
    resume = await generate_gemini_json(
        prompt=build_resume_prompt(job_description, current_resume),
        schema=RESUME_SCHEMA,
        empty_message="Gemini returned an empty response. Please try again.",
        invalid_message="Gemini returned an invalid resume response. Please try again.",
    )
    validate_resume(resume)
    return {"resume": resume}


@app.post("/api/application-pack")
async def generate_application_pack(payload: ResumeInputs) -> dict[str, Any]:
    job_description, current_resume = validate_inputs(
        payload,
        "Add both the current resume and job description before generating messages.",
    )
    ensure_gemini_key()
    application_pack = await generate_gemini_json(
        prompt=build_application_prompt(job_description, current_resume),
        schema=APPLICATION_PACK_SCHEMA,
        empty_message="Gemini returned an empty application pack. Please try again.",
        invalid_message="Gemini returned invalid application messages. Please try again.",
    )
    validate_application_pack(application_pack)
    return {"applicationPack": application_pack}


def extract_resume_text(content: bytes, suffix: str, content_type: str) -> str:
    if suffix == ".txt" or content_type == "text/plain":
        return content.decode("utf-8-sig", errors="replace")
    if suffix == ".docx":
        document = Document(BytesIO(content))
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)
    if suffix == ".pdf" or content_type == "application/pdf":
        reader = PdfReader(BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    raise ApiError("Upload a PDF, DOCX, or TXT resume.", 415)


def validate_inputs(payload: ResumeInputs, missing_message: str) -> tuple[str, str]:
    job_description = payload.jobDescription.strip()
    current_resume = payload.currentResume.strip()
    if not job_description or not current_resume:
        raise ApiError(missing_message, 400)
    if len(job_description) > MAX_TEXT_LENGTH or len(current_resume) > MAX_TEXT_LENGTH:
        raise ApiError("Each input must be 45,000 characters or fewer.", 413)
    return job_description, current_resume


def ensure_gemini_key() -> None:
    if not os.getenv("GEMINI_API_KEY"):
        raise ApiError(
            "The server is missing GEMINI_API_KEY. Add it to your .env file and restart the app.",
            503,
        )


async def generate_gemini_json(
    *, prompt: str, schema: dict[str, Any], empty_message: str, invalid_message: str
) -> dict[str, Any]:
    model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
    endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    request_body = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.2,
            "maxOutputTokens": 8192,
            "responseMimeType": "application/json",
            "responseSchema": schema,
        },
    }

    try:
        async with httpx.AsyncClient(timeout=90.0) as client:
            response = await client.post(
                endpoint,
                params={"key": os.environ["GEMINI_API_KEY"]},
                json=request_body,
            )
    except httpx.HTTPError as error:
        raise ApiError("Could not connect to Gemini. Please try again.", 502) from error

    try:
        data = response.json()
    except ValueError:
        data = {}

    if not response.is_success:
        provider_message = data.get("error", {}).get("message", "Gemini did not complete the request.")
        raise ApiError(provider_message, response.status_code)

    parts = data.get("candidates", [{}])[0].get("content", {}).get("parts", [])
    text = "".join(str(part.get("text", "")) for part in parts).strip()
    if not text:
        raise ApiError(empty_message, 502)

    try:
        parsed = json.loads(strip_code_fence(text))
    except (json.JSONDecodeError, TypeError) as error:
        raise ApiError(invalid_message, 502) from error
    if not isinstance(parsed, dict):
        raise ApiError(invalid_message, 502)
    return parsed


def strip_code_fence(value: str) -> str:
    value = re.sub(r"^```json\s*", "", value, flags=re.IGNORECASE)
    value = re.sub(r"^```\s*", "", value, flags=re.IGNORECASE)
    return re.sub(r"\s*```$", "", value).strip()


def validate_resume(resume: dict[str, Any]) -> None:
    required_arrays = ("education", "workExperience", "skills", "projects", "achievements")
    if not isinstance(resume.get("contact"), dict) or not isinstance(
        resume.get("professionalSummary"), str
    ):
        raise ApiError("Gemini returned an incomplete resume. Please try again.", 502)

    summary = resume["professionalSummary"].strip()
    if not summary or "\n" in summary or len(summary.split()) > 55:
        raise ApiError(
            "Gemini did not follow the one-line summary requirement. Please try again.", 502
        )
    if any(not isinstance(resume.get(key), list) for key in required_arrays):
        raise ApiError("Gemini returned an incomplete resume. Please try again.", 502)
    for job in resume["workExperience"]:
        if not isinstance(job, dict) or not isinstance(job.get("bullets"), list) or len(job["bullets"]) > 3:
            raise ApiError(
                "Gemini did not follow the three-bullet experience limit. Please try again.",
                502,
            )


def validate_application_pack(application_pack: dict[str, Any]) -> None:
    required_fields = ("coverLetter", "recruiterEmail", "linkedInDM", "referralRequest")
    if any(
        not isinstance(application_pack.get(field), str)
        or not application_pack[field].strip()
        for field in required_fields
    ):
        raise ApiError("Gemini returned incomplete application messages. Please try again.", 502)


# Vercel serves public/** from its CDN. Mount it only for local development.
if not os.getenv("VERCEL"):
    app.mount("/", StaticFiles(directory=PUBLIC_DIR, html=True), name="public")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "3000")))
